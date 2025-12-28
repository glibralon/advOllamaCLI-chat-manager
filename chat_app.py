# -*- coding: utf-8 -*-
"""
Updated command-line chat application with Ollama integration.

This version includes:
- A robust, interactive menu for session management.
- Dynamic model context length detection.
- In-chat file ingestion support for various formats, including spreadsheets.
- A new `/export --save` command with a unique delimiter-based extraction method.
- Support for exporting content to .docx, .md, .txt, .xlsx, and .pdf files.
- A user-friendly fallback if the export delimiters are not found.
- The ability to dynamically change the system prompt with a new in-chat command.
- Dynamic model switching from within a chat session.
- An 'Ollama Models' menu for listing and deleting local models.

This program offers a comprehensive set of features for managing conversations.

**Initial Setup & Command-Line Arguments:**
When starting the application, you can customize its behavior using the following arguments:
- `python chat_app.py --ingest <filepath>`: Ingests a file's content as a pre-prompt for the first message.

**Main Menu Features (Interactive):**
- **Start a New Chat:** Create a new conversation session.
- **Load an Existing Chat:** Resume a previously saved conversation from a JSON file.
- **Delete a Chat Session:** Permanently remove a saved chat history file.
- **Search Chat Sessions:** Find specific text or keywords across all saved chats.
- **Ollama Models:** Manage local Ollama models (list, delete, install).

**In-Chat Commands (While Interacting with the Model):**
- `/topic <prompt_text>`: Dynamically changes the chat session's topic.
- `/sys_prompt <prompt_text>`: Dynamically changes the system prompt for the current session.
- `/context <number>`: Adjusts the short-term memory window size (e.g., `/context 10` keeps the last 10 messages).
- `/summarize_context`: Manually triggers a conversation summary for long-term memory.
- `/ingest <filepath>`: Reads content from a file and includes it as part of your next prompt. Supports `.txt`, `.md`, `.docx`, `.pdf`, and spreadsheets.
- `/export --save <filename>`: Exports the **next assistant response** to a file. Supports `.md`, `.docx`, `.xlsx`, `.pdf`, and `.txt`. The filename and extension are optional, and the app will provide a default name and `.md` extension if none are given.
- `/export --save <filename> --all`: Exports the **entire chat history** to a file.
- `/model <model_name>`: Switch the current model being used for the chat.
- `/bye`: Ends the current chat session.

**Context Management:**
- The application automatically summarizes the conversation when it exceeds a set token limit to maintain long-term memory and prevent context window overflow.
- A short-term memory window of recent messages is maintained for the current session.
"""

import argparse
import json
import os
import re
import sys
import time
import tiktoken
import ollama
import textwrap
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple
import pandas as pd
from io import StringIO
from fpdf import FPDF
from halo import Halo
from docx import Document
from PyPDF2 import PdfReader

# --- Centralized Configuration Class ---
class AppConfig:
    """Manages application settings from a config file or defaults."""
    def __init__(self, config_file: str = "config.json"):
        self.config_file = config_file
        self.defaults: Dict[str, Any] = {
            "CONTEXT_WINDOW_SIZE": 8,
            "DEFAULT_SYSTEM_PROMPT": "You are a helpful consultant.",
            "DEFAULT_MODEL": 'neural-chat:7b',
            "CHATS_DIRECTORY": "chats",
            "AUTOMATIC_SUMMARY_THRESHOLD_RATIO": 0.75,
            "OLLAMA_HOST": "http://localhost:11434"
        }
        self.config = self.load_config()
        self.load_environment_overrides()

    def load_config(self) -> Dict[str, Any]:
        """Loads configuration from a JSON file."""
        try:
            with open(self.config_file, 'r') as f:
                return json.load(f)
        except FileNotFoundError:
            print(f"Warning: Configuration file '{self.config_file}' not found. Using default settings.")
            return {}
        except json.JSONDecodeError:
            print(f"Error: Invalid JSON in '{self.config_file}'. Using default settings.")
            return {}

    def load_environment_overrides(self):
        """Overrides settings with environment variables."""
        if os.getenv('OLLAMA_HOST'):
            self.config['OLLAMA_HOST'] = os.getenv('OLLAMA_HOST')
            
    def get(self, key: str, default: Optional[Any] = None) -> Any:
        """Retrieves a configuration value with a fallback."""
        return self.config.get(key, self.defaults.get(key, default))

# --- Ollama Service Class (Encapsulates API logic) ---
class OllamaService:
    """Handles all interactions with the local Ollama server."""
    def __init__(self, host: str):
        self.client = ollama.Client(host=host)
        
        # A dictionary of common model context lengths for robust fallback
        self.MODEL_CONTEXT_DEFAULTS = {
            "mixtral": 32768
        }
    
    def get_available_models(self) -> List[str]:
        """Fetches a list of available models from the local Ollama server."""
        try:
            response_obj = self.client.list()
            if not hasattr(response_obj, 'models'):
                print("The Ollama response object does not have a 'models' attribute.")
                return []
            models = response_obj.models
            if not models:
                print("Ollama returned an empty list of models.")
                return []
            model_names = []
            for model_obj in models:
                try:
                    if hasattr(model_obj, 'model'):
                        model_names.append(model_obj.model)
                except Exception as e:
                    print(f"An error occurred processing a single model entry: {e}")
            if not model_names:
                print("No valid models found in the response.")
                return []
            return model_names
        except Exception as e:
            print(f"An error occurred while fetching models.")
            print(f"Error details: {e}")
            return []

    def get_model_context_length(self, model_name: str) -> int:
        """
        Fetches the context length of a model using the Ollama API,
        with robust fallback mechanisms.
        """
        try:
            model_info = self.client.show(model_name)
            
            # FIRST PRIORITY: Access 'modelinfo' using dot notation, and then iterate its items.
            modelinfo_dict = getattr(model_info, 'modelinfo', {})
            for key, value in modelinfo_dict.items():
                if 'context_length' in key and isinstance(value, int):
                    return value
            
            # SECOND PRIORITY: Use regex to parse the 'modelfile' attribute, also with dot notation.
            modelfile_content = getattr(model_info, 'modelfile', '')
            num_ctx_match = re.search(r'num_ctx\s+(\d+)', modelfile_content)
            if num_ctx_match:
                return int(num_ctx_match.group(1))

            # THIRD PRIORITY: Check the hardcoded dictionary
            clean_model_name = model_name.split(':')[0]
            for key, length in self.MODEL_CONTEXT_DEFAULTS.items():
                if clean_model_name in key:
                    return length
            
            # FINAL FALLBACK
            return 4096
            
        except Exception as e:
            print(f"WARNING: Failed to determine context length for '{model_name}'. Using default. Error: {e}", file=sys.stderr)
            return 4096

    def chat(self, model: str, messages: List[Dict[str, str]], stream: bool = True, options: Optional[Dict[str, Any]] = None) -> Any:
        return self.client.chat(model=model, messages=messages, stream=stream, options=options)
    
    def generate_summary(self, conversation: List[Dict[str, Any]], model_name: str) -> str:
        """
        Generates a brief structured summary of the conversation using the specified Ollama model.
        
        Args:
            conversation (List[Dict[str, Any]]): The full conversation history.
            model_name (str): The name of the model to use for the summary.
            
        Returns:
            str: A JSON string containing the summary, or a fallback string if it fails.
        """
        summary_prompt_message = {
            "role": "user",
            "content": "Generate a concise JSON object. The object must contain two keys: 'narrative' and 'takeaways'. The 'narrative' value should be a brief paragraph summarizing the entire conversation. The 'takeaways' value should be a list of strings, with each string being a key takeaway or action item discussed. Do not include any other text, preambles, or explanations outside of the JSON object itself. Just the raw JSON."
        }
        
        # Send the entire conversation history followed by the new summary request
        messages_for_summary = conversation + [summary_prompt_message]
        
        try:
            response = self.chat(
                model=model_name,
                messages=messages_for_summary,
                stream=False,
                options={'temperature': 0.1}
            )
            summary_json_str = response['message']['content'].strip()
            
            if summary_json_str:
                return summary_json_str
            else:
                print("Warning: Received an empty summary from the model.", file=sys.stderr)
                return "No summary available."
        except ollama.OllamaAPIError as e:
            print(f"Ollama API Error during summary generation: {e}", file=sys.stderr)
            return "No summary available."
        except Exception as e:
            print(f"Warning: Failed to generate summary. General error: {e}", file=sys.stderr)
            return "No summary available."

# --- Helper Functions ---
try:
    tokenizer = tiktoken.get_encoding("cl100k_base")
except Exception as e:
    print(f"Warning: Failed to load tokenizer for token counting. Context management will be based on message count only. Error: {e}", file=sys.stderr)
    tokenizer = None

# --- Delimiter Constants for Robust Export ---
BLOCK_START = "<<<BLOCK_START>>>"
BLOCK_END = "<<<BLOCK_END>>>"

def count_tokens(text: str) -> int:
    """Counts tokens in a string using a tokenizer."""
    if tokenizer:
        return len(tokenizer.encode(text))
    return len(text.split())

def sanitize_filename(name: str) -> str:
    """Sanitizes a string to be a valid filename."""
    return re.sub(r'[\\/:*?"<>|]', '_', name)


def truncate_summary_for_display(summary_text: str, max_lines: int = 5) -> str:
    """
    Truncates a multi-line or long string to a maximum number of lines for display.
    If the text contains no newlines, it will wrap the text to simulate lines.
    
    Args:
        summary_text (str): The summary text to truncate.
        max_lines (int): The maximum number of lines to display.
        
    Returns:
        str: The truncated summary text.
    """
    if '\n' in summary_text:
        lines = summary_text.strip().split('\n')
    else:
        lines = textwrap.wrap(summary_text, width=100)
    
    if len(lines) > max_lines:
        truncated_text = "\n".join(lines[:max_lines])
        return truncated_text + "..."
    
    return summary_text

def ingest_file_content(filepath: str) -> Optional[str]:
    """
    Ingests content from a file, converting spreadsheet data to a string representation.
    
    Args:
        filepath (str): The path to the file to ingest.
    
    Returns:
        Optional[str]: The file content as a string, or None if an error occurred.
    """
    if not os.path.exists(filepath):
        print(f"Error: File not found at {filepath}", file=sys.stderr)
        return None

    _, file_extension = os.path.splitext(filepath)
    file_extension = file_extension.lower()
    content = ""

    try:
        if file_extension in ['.txt', '.py', '.json', '.csv', '.md']:
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read()
        elif file_extension == '.docx':
            doc = Document(filepath)
            full_text = [para.text for para in doc.paragraphs]
            content = '\n'.join(full_text)
        elif file_extension == '.pdf':
            reader = PdfReader(filepath)
            full_text = [page.extract_text() for page in reader.pages]
            content = '\n'.join(full_text)
        elif file_extension in ['.xls', '.xlsx']:
            df = pd.read_excel(filepath)
            return f"Context from spreadsheet '{os.path.basename(filepath)}':\n\n{df.to_markdown(index=False)}"
        else:
            print(f"Error: Unsupported file type '{file_extension}'.", file=sys.stderr)
            return None
    except Exception as e:
        print(f"Error ingesting file: {e}", file=sys.stderr)
        return None

def extract_content_from_block(text: str) -> Optional[str]:
    """
    Extracts content between the BLOCK_START and BLOCK_END delimiters.
    
    Args:
        text (str): The text to search within.
        
    Returns:
        Optional[str]: The extracted content, or None if delimiters are not found.
    """
    pattern = re.compile(f'{re.escape(BLOCK_START)}(.*?){re.escape(BLOCK_END)}', re.DOTALL)
    match = pattern.search(text)
    if match:
        return match.group(1).strip()
    return None

def remove_think_tags(text: str) -> str:
    """
    Removes any content between <think> and </think> tags.
    
    Args:
        text (str): The text to process.
        
    Returns:
        str: The text with the tags and their content removed.
    """
    return re.sub(r'<think>.*?</think>', '', text, flags=re.DOTALL).strip()


def save_excel_file(content: str, filename: str):
    """
    Parses a Markdown table from the content and saves it to an Excel file.
    
    Args:
        content (str): The Markdown table content.
        filename (str): The name of the file to save.
    """
    try:
        # Find the start and end of the Markdown table
        table_start = content.find('|')
        table_end = content.rfind('|')
        
        if table_start == -1 or table_end == -1:
            raise ValueError("No Markdown table found in the response.")
        
        table_string = content[table_start:].strip()
        data = StringIO(table_string)
        
        df = pd.read_csv(data, sep='|', engine='python', skipinitialspace=True)
        df.columns = df.columns.str.strip()
        df = df.iloc[1:].reset_index(drop=True)
        df.to_excel(filename, index=False)
        
        print(f"✅ Content successfully saved to {filename}")
    except Exception as e:
        print(f"Error saving to .xlsx: {e}")
        raise

def save_pdf_file(content: str, filename: str):
    """
    Saves Markdown-formatted content to a PDF file with flexible font handling.
    
    Args:
        content (str): The Markdown content to save.
        filename (str): The name of the file to save.
    """
    try:
        pdf = FPDF()
        font_name = 'DejaVu'
        regular_font_path = 'DejaVuSans.ttf'
        bold_font_path = 'DejaVuSans-Bold.ttf'
        font_found = False
        if os.path.exists(regular_font_path) and os.path.exists(bold_font_path):
            try:
                pdf.add_font(font_name, '', regular_font_path, uni=True)
                pdf.add_font(font_name, 'B', bold_font_path, uni=True)
                font_found = True
            except RuntimeError as e:
                print(f"Warning: Failed to load DejaVu font: {e}")

        if font_found:
            pdf.set_font(font_name, size=12)
        else:
            print("Warning: DejaVuSans.ttf not found. Using 'Helvetica' as a fallback. "
                  "Unicode characters may not render correctly.")
            pdf.set_font('Helvetica', size=12)

        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        
        paragraphs = content.split('\n')
        for para in paragraphs:
            para = para.strip()
            if para.startswith('# '):
                pdf.set_font(style='B', size=20)
                pdf.multi_cell(0, 10, txt=para[2:])
                pdf.ln(2)
                if font_found:
                    pdf.set_font(font_name, size=12)
                else:
                    pdf.set_font('Helvetica', size=12)
            elif para.startswith('## '):
                pdf.set_font(style='B', size=16)
                pdf.multi_cell(0, 8, txt=para[3:])
                pdf.ln(2)
                if font_found:
                    pdf.set_font(font_name, size=12)
                else:
                    pdf.set_font('Helvetica', size=12)
            elif para:
                pdf.multi_cell(0, 6, txt=para)
                pdf.ln()

        pdf.output(filename)
        print(f"✅ Content successfully saved to {filename}")

    except Exception as e:
        print(f"Error saving to .pdf: {e}")
        
def save_content_to_file(content: str, filename: str, file_type: str):
    """Saves content to a file based on the specified file_type."""
    if file_type == 'docx':
        try:
            document = Document()
            document.add_paragraph(content)
            document.save(filename)
            print(f"✅ Content successfully saved to {filename}")
        except Exception as e:
            print(f"Error saving to .docx: {e}")
    elif file_type == 'md':
        try:
            with open(filename, 'w', encoding='utf-8') as f:
                f.write(content)
            print(f"✅ Content successfully saved to {filename}")
        except Exception as e:
            print(f"Error saving to .md: {e}")
    elif file_type in ['excel', 'xlsx']:
        try:
            save_excel_file(content, filename)
        except Exception as e:
            print(f"Error saving to .xlsx: {e}")
    elif file_type == 'pdf':
        try:
            save_pdf_file(content, filename)
        except Exception as e:
            print(f"Error saving to .pdf: {e}")
    else: # Default to plain text
        try:
            with open(filename, 'w', encoding='utf-8') as f:
                f.write(content)
            print(f"✅ Content successfully saved to {filename}")
        except Exception as e:
            print(f"Error saving to .txt: {e}")

def handle_special_commands(user_input: str, session_manager: 'SessionManager') -> Tuple[Optional[str], Optional[str], bool]:
    """
    Handles special commands and returns the modified input, a special flag, and the --all flag state.
    
    Args:
        user_input (str): The user's raw input.
        session_manager (SessionManager): The current session manager instance.
        
    Returns:
        Tuple[Optional[str], Optional[str], bool]: A tuple containing the processed prompt text, a special flag for
        command type ('save', 'bye', or None), and a boolean for the export '--all' flag.
    """
    export_pattern = re.compile(r'/export\s+--save(?:\s+(".*?"|\S+))?(?:\s+--all)?', re.IGNORECASE)
    export_match = export_pattern.search(user_input)
    
    if export_match:
        print("Export command detected.")
        prompt_text = export_pattern.sub('', user_input).strip()
        filename_match = export_match.group(1)
        all_flag = '--all' in export_match.group(0).lower()
        session_manager.export_filename = filename_match.strip('\"') if filename_match else None
        
        if session_manager.export_filename:
           base_name, file_extension = os.path.splitext(session_manager.export_filename)
           if not file_extension:
               if all_flag:
                   session_manager.export_filename = f"{base_name}.md"
                   print(f"No file extension provided with --all flag. Defaulting to: {session_manager.export_filename}")
               else:
                   session_manager.export_filename = f"{base_name}.txt"
                   print(f"No file extension provided. Defaulting to: {session_manager.export_filename}")
       
        if session_manager.export_filename:
            file_extension = os.path.splitext(session_manager.export_filename)[1].lower()
            if file_extension in ['.doc', '.docx']:
                session_manager.export_type = 'docx'
                if file_extension == '.doc':
                    base_name, _ = os.path.splitext(session_manager.export_filename)
                    session_manager.export_filename = f"{base_name}.docx"
            elif file_extension == '.md':
                session_manager.export_type = 'md'
            elif file_extension in ['.xls', '.xlsx']:
                session_manager.export_type = 'excel'
                if file_extension == '.xls':
                    base_name, _ = os.path.splitext(session_manager.export_filename)
                    session_manager.export_filename = f"{base_name}.xlsx"
            elif file_extension == '.pdf':
                session_manager.export_type = 'pdf'
            elif file_extension in ['.py', '.js', '.html', '.cpp']:
                session_manager.export_type = 'txt'
            else:
                session_manager.export_type = file_extension.lstrip('.')
        else:
            session_manager.export_type = 'md'

        return prompt_text, 'save', all_flag
    
    model_match = re.match(r'/model\s+(\S+)', user_input.strip())
    if model_match:
        new_model = model_match.group(1).strip()
        available_models = session_manager.ollama_service.get_available_models()
        if new_model in available_models:
            old_model = session_manager.model_name
            session_manager.set_model(new_model)
            command_message = {
                "role": "command", 
                "content": f"Model switched from '{old_model}' to '{new_model}'.", 
                "timestamp": datetime.now().isoformat()
            }
            session_manager.conversation_history.append(command_message)
            print(f"✨ Model for this session has been updated to: '{new_model}'.")
            print(f"✨ New model context length is: {session_manager.model_context_length} tokens.")
        else:
            print(f"Error: Model '{new_model}' not found. Please ensure it is installed locally.")
        return None, None, False

    topic_match = re.match(r'/topic\s+(.*)', user_input.strip(), re.DOTALL)
    if topic_match:
        new_topic = topic_match.group(1).strip()
        if new_topic:
            session_manager.set_topic_name(new_topic)
            command_message = {
                "role": "command", 
                "content": f"Session topic updated to: '{new_topic}'", 
                "timestamp": datetime.now().isoformat()
            }
            session_manager.conversation_history.append(command_message)
            print(f"✨ Session topic has been updated to: '{new_topic}'")
        else:
            print("Error: Topic name cannot be empty. Usage: /topic <new topic text>")
        return None, None, False
    
    context_match = re.match(r'/context\s+(\d+)', user_input.strip())
    if context_match:
        try:
            new_size = int(context_match.group(1))
            if new_size > 0:
                session_manager.context_window_size = new_size
                print(f"✨ Context window size for this session is now set to {session_manager.context_window_size} messages.")
            else:
                print("Error: Context window size must be a positive number.")
        except ValueError:
            print("Error: Invalid number for context window size. Usage: /context <number>")
        return None, None, False
        
    sys_prompt_match = re.match(r'/sys_prompt\s+(.*)', user_input.strip(), re.DOTALL)
    if sys_prompt_match:
        new_prompt = sys_prompt_match.group(1).strip()
        if new_prompt:
            session_manager.set_system_prompt(new_prompt)
            command_message = {
                "role": "command", 
                "content": f"System prompt updated to: '{new_prompt}'", 
                "timestamp": datetime.now().isoformat()
            }
            session_manager.conversation_history.append(command_message)
            print(f"✨ System prompt for this session has been updated to: '{new_prompt}'")
        else:
            print("Error: System prompt cannot be empty. Usage: /sys_prompt <new prompt text>")
        return None, None, False

    if user_input.strip().lower() == '/summarize_context':
        if len(session_manager.conversation_history) > 0:
            print("Generating a summary of the entire conversation for long-term memory...")
            spinner = Halo(text='Generating summary...', spinner='dots')
            spinner.start()
            try:
                full_conversation = [msg for msg in session_manager.conversation_history if msg['role'] in ['user', 'assistant']]
                summary_json_str = session_manager.ollama_service.generate_summary(full_conversation, session_manager.model_name)
                
                try:
                    summary_data = json.loads(summary_json_str)
                    session_manager.conversation_history = [msg for msg in session_manager.conversation_history if msg['role'] != 'summary']
                    
                    if summary_data.get("narrative") and summary_data.get("takeaways"):
                        narrative = summary_data["narrative"]
                        takeaways_str = "\n".join([f"• {t}" for t in summary_data["takeaways"]])
                        summary_display_text = f"{narrative}\n\nKey Takeaways:\n{takeaways_str}"
                        
                        summary_message = {
                            "role": "summary", 
                            "content": summary_display_text, 
                            "timestamp": datetime.now().isoformat()
                        }
                        session_manager.conversation_history.append(summary_message)
                        session_manager.chat_data["summary"] = summary_data
                        spinner.succeed("✅ Automatic summary added for long-term memory.")
                    else:
                        spinner.fail("❌ Failed to generate a valid summary. No summary added.")
                except json.JSONDecodeError:
                    print("Warning: Failed to parse structured summary. Saving as a plain string.")
                    session_manager.chat_data["summary"] = {"narrative": summary_json_str, "takeaways": []}
                    summary_message = {
                        "role": "summary", 
                        "content": summary_json_str, 
                        "timestamp": datetime.now().isoformat()
                    }
                    session_manager.conversation_history.append(summary_message)
                    spinner.succeed("✅ Summary generated and added as a plain string.")
            except Exception as e:
                spinner.fail(f"Error generating summary: {e}")
        else:
            print("Cannot summarize an empty conversation.")
        return None, None, False
        
    ingest_pattern = re.compile(r'/ingest\s+(".*?"|\S+)')
    matches = ingest_pattern.findall(user_input)
    if matches:
        modified_input = user_input
        for match in matches:
            filepath = match.strip('\"')
            ingested_text = ingest_file_content(filepath)
            if ingested_text:
                modified_input = modified_input.replace(f"/ingest {match}", f"### Content from '{os.path.basename(filepath)}':\n\n{ingested_text}\n\n", 1)
            else:
                print(f"File ingestion failed for '{filepath}'. The prompt will not be sent.")
                return None, None, False
        return modified_input, None, False
    
    if user_input.strip().lower() == '/bye':
        return None, 'bye', False
    
    processed_input = remove_think_tags(user_input)
    return processed_input, None, False

# --- Session Manager Class ---
class SessionManager:
    """Manages the lifecycle of a single chat session."""
    def __init__(self, ollama_service: OllamaService, app_config: AppConfig, model_name: str, system_prompt: str, session_file: Optional[str] = None, new_topic: Optional[str] = None):
        os.makedirs(app_config.get("CHATS_DIRECTORY"), exist_ok=True)
        self.ollama_service = ollama_service
        self.app_config = app_config
        self.model_name = model_name
        self.clean_model_name = sanitize_filename(model_name).replace(':', '_')
        self.chat_data: Dict[str, Any] = {}
        self.history_file_path: str = ""
        self.export_filename: Optional[str] = None
        self.export_type: Optional[str] = None
        self.system_prompt = system_prompt
        self.context_window_size = self.app_config.get("CONTEXT_WINDOW_SIZE")
    
        self.model_context_length = self.ollama_service.get_model_context_length(self.model_name)
        self.automatic_summary_threshold = int(self.model_context_length * self.app_config.get("AUTOMATIC_SUMMARY_THRESHOLD_RATIO"))

        if session_file:
            self.load_session(session_file)
        else:
            self.start_new_session(new_topic)
    
    def set_model(self, new_model: str) -> None:
        """Updates the session's model and related attributes."""
        self.chat_data["model_name"] = new_model
        old_model = self.model_name
        self.model_name = new_model
        self.clean_model_name = sanitize_filename(new_model).replace(':', '_')
        self.model_context_length = self.ollama_service.get_model_context_length(self.model_name)
        self.automatic_summary_threshold = int(self.model_context_length * self.app_config.get("AUTOMATIC_SUMMARY_THRESHOLD_RATIO"))
        
        if self.history_file_path:
            self.update_filename_for_topic()
            
    def load_session(self, session_file: str) -> None:
        file_path = os.path.join(self.app_config.get("CHATS_DIRECTORY"), session_file)
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                self.chat_data = json.load(f)
            self.history_file_path = file_path
            
            summary = self.chat_data.get('summary', 'No summary available.')
            if isinstance(summary, dict):
                narrative = summary.get("narrative", "")
                takeaways = summary.get("takeaways", [])
                takeaways_str = "\n".join([f"• {t}" for t in takeaways])
                summary_display_text = f"{narrative}\n\nKey Takeaways:\n{takeaways_str}"
            else:
                summary_display_text = summary

            print(f"Loaded existing chat session '{self.chat_data.get('topic_name', 'Unnamed')}' from '{file_path}'.")
            print(f"Summary: {summary_display_text}")
            
            self.system_prompt = self.chat_data.get("system_prompt", self.app_config.get('DEFAULT_SYSTEM_PROMPT'))

            conversation_msgs = [msg for msg in self.chat_data.get('conversation', []) if msg['role'] in ['user', 'assistant']]
            if conversation_msgs:
                print(f"\n--- Last {min(self.context_window_size, len(conversation_msgs))} messages ---")
                for msg in conversation_msgs[-self.context_window_size:]:
                    print(f"[{msg['role'].upper()}]:\n{msg['content']}\n")
        except (FileNotFoundError, json.JSONDecodeError) as e:
            print(f"Error loading session from '{file_path}': {e}. Starting a new session.", file=sys.stderr)
            self.start_new_session()

    def start_new_session(self, new_topic: Optional[str] = None) -> None:
        self.chat_data = {
            "model_name": self.model_name,
            "topic_name": new_topic if new_topic else "New Session",
            "system_prompt": self.system_prompt,
            "conversation": [],
            "total_response_time_seconds": 0.0,
            "total_ttft_seconds": 0.0,
            "creation_date": datetime.now().isoformat()
        }
        self.set_topic_name(new_topic if new_topic else "New Session")
        print("Starting a new chat session.")
    
    def _get_unique_filename(self, topic: str) -> str:
        """Generates a unique filename for the chat session."""
        sanitized_topic = sanitize_filename(topic)
        base_filename = f"chat_history_{self.clean_model_name}_{sanitized_topic}"
        file_path = os.path.join(self.app_config.get("CHATS_DIRECTORY"), f"{base_filename}.json")
        
        counter = 1
        while os.path.exists(file_path):
            file_path = os.path.join(self.app_config.get("CHATS_DIRECTORY"), f"{base_filename}_{str(counter).zfill(2)}.json")
            counter += 1
        return file_path

    def save_session(self) -> None:
        """Saves the current session data to its history file."""
        full_conversation = [msg for msg in self.chat_data.get("conversation", []) if msg['role'] in ['user', 'assistant']]
        
        current_summary = self.chat_data.get("summary")
        
        if not current_summary or current_summary == "No summary available.":
            if full_conversation:
                print("Generating session summary before saving...")
                spinner = Halo(text='Generating summary...', spinner='dots')
                spinner.start()
                try:
                    summary_json_str = self.ollama_service.generate_summary(full_conversation, self.model_name)
                   
                    summary_json_str = re.sub(r'```json|```', '', summary_json_str).strip()
                    try:
                        summary_data = json.loads(summary_json_str)
                        self.chat_data["summary"] = summary_data
                        spinner.succeed("✅ Generated a structured summary.")
                    except json.JSONDecodeError:
                        print("Warning: Failed to parse structured summary. Saving as a plain string.")
                        self.chat_data["summary"] = {"narrative": summary_json_str, "takeaways": []}
                        spinner.succeed("✅ Summary generated and added as a plain string.")
                except Exception as e:
                    spinner.fail(f"Error generating summary: {e}")
                    self.chat_data["summary"] = {"narrative": "No summary available.", "takeaways": []}
            else:
                self.chat_data["summary"] = {"narrative": "No conversation to summarize.", "takeaways": []}
        
        self.update_filename_for_topic()
        with open(self.history_file_path, 'w', encoding='utf-8') as f:
            json.dump(self.chat_data, f, indent=4, ensure_ascii=False)
        print(f"Chat history saved to '{self.history_file_path}'.")

    def update_filename_for_topic(self, new_topic_name: Optional[str] = None) -> None:
        """
        Updates the session's filename based on the current model and topic name.
        This function now renames the existing file instead of creating a new one.
        
        Args:
            new_topic_name (Optional[str]): The new topic name, if provided.
        """
        old_file_path = self.history_file_path
        
        topic_for_new_name = new_topic_name if new_topic_name else self.chat_data.get("topic_name", "New Session")
        
        sanitized_topic_for_new_name = sanitize_filename(topic_for_new_name)
        new_base_filename = f"chat_history_{self.clean_model_name}_{sanitized_topic_for_new_name}"
        
        new_file_path = os.path.join(self.app_config.get("CHATS_DIRECTORY"), f"{new_base_filename}.json")

        if os.path.exists(old_file_path) and old_file_path != new_file_path:
            counter = 1
            unique_new_file_path = new_file_path
            while os.path.exists(unique_new_file_path):
                unique_new_file_path = os.path.join(
                    self.app_config.get("CHATS_DIRECTORY"),
                    f"{new_base_filename}_{str(counter).zfill(2)}.json"
                )
                counter += 1
            
            os.rename(old_file_path, unique_new_file_path)
            self.history_file_path = unique_new_file_path
            print(f"Session file renamed to '{os.path.basename(self.history_file_path)}'.")
        elif not os.path.exists(old_file_path):
            self.history_file_path = self._get_unique_filename(topic_for_new_name)
    
    @property
    def conversation_history(self) -> List[Dict[str, Any]]:
        return self.chat_data["conversation"]

    @property
    def topic_name(self) -> str:
        return self.chat_data.get("topic_name", "Unnamed Session")

    def set_topic_name(self, name: str) -> None:
        """Updates the session's topic and renames the file."""
        self.chat_data["topic_name"] = name
        self.update_filename_for_topic()

    def set_system_prompt(self, prompt: str) -> None:
        self.system_prompt = prompt
        self.chat_data["system_prompt"] = prompt
        
def list_sessions(chats_directory: str) -> List[str]:
    """
    Lists all chat session files in the chats directory, showing a truncated summary.
    
    Args:
        chats_directory (str): The directory where chat sessions are stored.
        
    Returns:
        List[str]: A list of filenames for the available sessions.
    """
    files = [f for f in os.listdir(chats_directory) if f.endswith('.json')]
    if not files:
        print("No saved chat sessions found.")
    else:
        print("\n--- Available Chat Sessions ---")
        for i, file in enumerate(files):
            try:
                with open(os.path.join(chats_directory, file), 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    topic = data.get('topic_name', os.path.splitext(file)[0])
                    model = data.get('model_name', 'Unknown Model')
                    summary_data = data.get('summary', 'No summary available.')
                    
                    summary_to_display = ""
                    if isinstance(summary_data, dict) and 'narrative' in summary_data:
                        summary_to_display = truncate_summary_for_display(summary_data['narrative'], 3)
                    elif isinstance(summary_data, str):
                        summary_to_display = truncate_summary_for_display(summary_data, 3)
                    else:
                        summary_to_display = "No summary available."

                    print(f"[{i + 1}] Topic: '{topic}' | Model: {model}")
                    print(f"    Summary: {summary_to_display}\n")
                    
            except (json.JSONDecodeError, KeyError):
                print(f"[{i + 1}] Corrupted or unknown session file: {file}")
    print("-------------------------------")
    return files

def search_sessions(query_string: str, chats_directory: str) -> None:
    """
    Searches for a query string in chat session files with enhanced features.
    
    Args:
        query_string (str): The search query, which may include flags.
        chats_directory (str): The directory to search within.
    """
    query_parts = query_string.split()
    query = []
    case_sensitive = False
    role_filter = None
    
    i = 0
    while i < len(query_parts):
        part = query_parts[i]
        if part.lower() == '--case-sensitive':
            case_sensitive = True
        elif part.lower() == '--role' and i + 1 < len(query_parts):
            role_filter = query_parts[i+1].lower()
            if role_filter not in ['user', 'assistant']:
                print(f"Warning: Invalid role '{role_filter}'. Valid roles are 'user' or 'assistant'. Ignoring role filter.")
                role_filter = None
            i += 1
        else:
            query.append(part)
        i += 1
        
    final_query = " ".join(query).strip()

    if not final_query:
        print("Search query cannot be empty.")
        return

    print(f"\n--- Searching for '{final_query}' in chat histories... ---")

    regex_flags = 0 if case_sensitive else re.IGNORECASE
    
    if ' ' in final_query:
        pattern = re.compile(re.escape(final_query), flags=regex_flags)
    else:
        pattern = re.compile(r'\b' + re.escape(final_query) + r'\b', flags=regex_flags)
    
    found_results = False
    files = [f for f in os.listdir(chats_directory) if f.endswith('.json')]
    
    if not files:
        print("No chat sessions found to search.")
        return

    for filename in files:
        file_path = os.path.join(chats_directory, filename)
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                chat_data = json.load(f)
            
            conversation = chat_data.get('conversation', [])
            topic_name = chat_data.get('topic_name', 'Unnamed Session')
            
            matches_in_file = []
            for message in conversation:
                if role_filter and message.get('role', '').lower() != role_filter:
                    continue

                content = message.get('content', '')
                for match in pattern.finditer(content):
                    matches_in_file.append({
                        'role': message['role'],
                        'timestamp': message.get('timestamp', 'N/A'),
                        'content': content,
                        'match_text': match.group(0),
                        'match_start': match.start(),
                        'match_end': match.end()
                    })
            
            if matches_in_file:
                found_results = True
                print(f"\n📂 File: {filename} | Topic: '{topic_name}'")
                
                for match in matches_in_file:
                    content_to_display = match['content']
                    
                    lines = content_to_display.split('\n')
                    match_line_num = 0
                    char_count = 0
                    
                    for i, line in enumerate(lines):
                        char_count += len(line) + 1
                        if char_count > match['match_start']:
                            match_line_num = i
                            break
                    
                    start_line = max(0, match_line_num - 2)
                    end_line = min(len(lines), match_line_num + 3)
                    
                    snippet_lines = lines[start_line:end_line]
                    highlighted_snippet_lines = []
                    
                    for i, line in enumerate(snippet_lines):
                        if start_line + i == match_line_num:
                            highlighted_line = re.sub(
                                pattern, 
                                lambda m: f"\033[1m{m.group(0)}\033[0m", 
                                line
                            )
                            highlighted_snippet_lines.append(highlighted_line)
                        else:
                            highlighted_snippet_lines.append(line)
                            
                    snippet_with_context = '\n'.join(highlighted_snippet_lines)
                    
                    print(f"  > [{match['role'].upper()}] ({match.get('timestamp', 'N/A')}) - ...\n{snippet_with_context}\n...")
        
        except (json.JSONDecodeError, KeyError) as e:
            print(f"Warning: Could not read or parse file {filename}. Error: {e}", file=sys.stderr)
    
    if not found_results:
        print("No matching content found.")
    print("--------------------------------------------------")

def run_chat_session(ollama_service: OllamaService, session_manager: 'SessionManager', app_config: AppConfig, initial_ingest_file: Optional[str]):
    """The main chat loop for a single session."""
    
    print(f"\n--- Chat Session Started, see previous messages above ---")
    print(f"Current Model: {session_manager.model_name}")
    print(f"Topic: '{session_manager.topic_name}'")
    print("---------------------------\n")
    if initial_ingest_file:
        ingested_text = ingest_file_content(initial_ingest_file)
        if ingested_text:
            print(f"Pre-prompting model with content from {initial_ingest_file}...")
            initial_prompt_content = f"Context from file '{os.path.basename(initial_ingest_file)}':\n{ingested_text}"
            initial_prompt_tokens = count_tokens(initial_prompt_content)
            session_manager.conversation_history.insert(0, {
                'role': 'user', 
                'content': initial_prompt_content, 
                'token_count': initial_prompt_tokens,
                'timestamp': datetime.now().isoformat()
            })
        else:
            print(f"Warning: Failed to ingest file specified by --ingest. Continuing without it.")

    total_response_time = session_manager.chat_data.get("total_response_time_seconds", 0.0)
    total_ttft = session_manager.chat_data.get("total_ttft_seconds", 0.0)

    print(f"\nUse '/topic <text>' to change the session topic.")
    print(f"Use '/sys_prompt <text>' to change the system prompt.")
    print(f"Use '/ingest <filepath>' to add a file to the conversation.")
    print(f"Use '/export --save <filename>' to save the next response to a file (e.g., .md, .docx, .xlsx).")
    print(f"Use '/export --save <filename> --all' to save the entire chat history.")
    print(f"Use '/model <model_name>' to change the model for this session.")
    print(f"Use '/summarize_context' to manually generate a long-term memory summary.")
    print(f"\n--- Session Configuration ---")
    print(f"System Prompt: '{session_manager.system_prompt}'")
    print(f"Current context window size: {session_manager.context_window_size} messages. Use '/context <number>' to change it.")
    print(f"Model context window length: {session_manager.model_context_length} tokens.")
    print(f"Automatic summary threshold: {session_manager.automatic_summary_threshold} tokens.")
    
    print(f"\nType 'EOF' on a new line to send your prompt or '/bye' to exit.")

    try:
        while True:
            print("\n>>> ", end='', flush=True)
            lines = []
            while True:
                line = input()
                if line.strip().lower() == '/bye':
                    user_input = '/bye'
                    break
                if line.strip().lower() == 'eof':
                    user_input = "\n".join(lines)
                    break
                lines.append(line)
            
            processed_input, special_prompt_flag, export_all = handle_special_commands(user_input, session_manager)
            
            if special_prompt_flag == 'bye':
                print("Ending chat.")
                break

            if processed_input is None:
                continue
            
            if not export_all:
                if special_prompt_flag == 'save':
                    if session_manager.export_type in ['excel', 'xlsx']:
                        excel_instructions = f"Generate your response as a single, valid Markdown table with headers. Enclose the table within the delimiters {BLOCK_START} and {BLOCK_END}. Do not include any other text, code blocks, or explanations before or after the table. Only provide the table itself."
                        processed_input = f"{processed_input}\n\n{excel_instructions}"
                    else:
                        markdown_instructions = f"Please use Markdown for structured output. Enclose the entire content to be saved within the delimiters {BLOCK_START} and {BLOCK_END}. Do not include any other text."
                        processed_input = f"{processed_input}\n\n{markdown_instructions}"

                user_message_obj = {
                    "role": "user", 
                    "content": processed_input, 
                    "token_count": count_tokens(processed_input),
                    "timestamp": datetime.now().isoformat()
                }
                session_manager.conversation_history.append(user_message_obj)

                context_messages = [{"role": "system", "content": session_manager.system_prompt}]
                 
                last_summary_index = next((i for i, msg in reversed(list(enumerate(session_manager.conversation_history))) if msg['role'] == 'summary'), -1)
                
                if last_summary_index == -1:
                    messages_to_consider = [
                        msg for msg in session_manager.conversation_history if msg['role'] in ['user', 'assistant']
                    ]
                else:
                    messages_to_consider = [
                        msg for msg in session_manager.conversation_history[last_summary_index + 1:] if msg['role'] in ['user', 'assistant']
                    ]
                
                tokens_in_new_messages = count_tokens(
                    " ".join([msg['content'] for msg in messages_to_consider])
                )
                
                if tokens_in_new_messages > session_manager.automatic_summary_threshold and tokenizer:
                    print("ℹ️ Conversation history is growing. Generating automatic summary for context...")
                    spinner = Halo(text='Generating summary...', spinner='dots')
                    spinner.start()
                    try:
                        full_conversation_for_summary = [msg for msg in session_manager.conversation_history if msg['role'] in ['user', 'assistant']]
                        summary_json_str = session_manager.ollama_service.generate_summary(full_conversation_for_summary, session_manager.model_name)
                        
                        summary_json_str = re.sub(r'```json|```', '', summary_json_str).strip()
                        
                        try:
                            summary_data = json.loads(summary_json_str)
                            session_manager.conversation_history = [
                                msg for msg in session_manager.conversation_history if msg['role'] != 'summary'
                            ]
                            if summary_data.get("narrative") and summary_data.get("takeaways"):
                                narrative = summary_data["narrative"]
                                takeaways_str = "\n".join([f"• {t}" for t in summary_data["takeaways"]])
                                summary_display_text = f"{narrative}\n\nKey Takeaways:\n{takeaways_str}"

                                summary_message = {
                                    "role": "summary", 
                                    "content": summary_display_text, 
                                    "timestamp": datetime.now().isoformat()
                                }
                                session_manager.conversation_history.append(summary_message)
                                session_manager.chat_data["summary"] = summary_data
                                spinner.succeed("✅ Automatic summary added for long-term memory.")
                            else:
                                spinner.fail("❌ Failed to generate a valid summary. No summary added.")
                        except json.JSONDecodeError:
                            print("Warning: Failed to parse structured summary. Saving as a plain string.")
                            session_manager.chat_data["summary"] = {"narrative": summary_json_str, "takeaways": []}
                            summary_message = {
                                "role": "summary", 
                                "content": summary_json_str, 
                                "timestamp": datetime.now().isoformat()
                            }
                            session_manager.conversation_history.append(summary_message)
                            spinner.succeed("✅ Summary generated and added as a plain string.")
                    except Exception as e:
                        spinner.fail(f"Error generating summary: {e}")

                summary_message = next((msg for msg in reversed(session_manager.conversation_history) if msg['role'] == 'summary'), None)
                
                if summary_message:
                    context_messages.append(summary_message)
                    
                recent_messages = [msg for msg in session_manager.conversation_history if msg['role'] in ['user', 'assistant']][-session_manager.context_window_size:]
                context_messages.extend(recent_messages)
                
                start_total_time = time.time()
                spinner = Halo(spinner='dots')
                
                try:
                    spinner.start()
                    stream = session_manager.ollama_service.chat(model=session_manager.model_name, messages=context_messages, stream=True)

                    assistant_response_content = ""
                    first_chunk = True
                    time_to_first_token = 0.0
                    
                    for chunk in stream:
                        if first_chunk:
                            end_ttft_time = time.time()
                            time_to_first_token = end_ttft_time - start_total_time
                            first_chunk = False
                            if hasattr(spinner, '_spinner_thread') and spinner._spinner_thread:
                                spinner.stop_and_persist()
                                print()
                        content_chunk = chunk['message']['content']
                        print(content_chunk, end='', flush=True)
                        assistant_response_content += content_chunk
                    
                    print()

                except KeyboardInterrupt:
                    if hasattr(spinner, '_spinner_thread') and spinner._spinner_thread:
                        spinner.stop()
                    print("\n\nResponse interrupted by user. Continuing chat. 🛑")
                    session_manager.conversation_history.pop()
                    continue

                end_total_time = time.time()
                total_time_for_turn = end_total_time - start_total_time

                total_response_time += total_time_for_turn
                total_ttft += time_to_first_token

                clean_assistant_response = remove_think_tags(assistant_response_content)
                
                assistant_message_obj = {
                    "role": "assistant",
                    "content": clean_assistant_response,
                    "model_used": session_manager.model_name,
                    "token_count": count_tokens(clean_assistant_response),
                    "response_time_seconds": round(total_time_for_turn, 2),
                    "time_to_first_token_seconds": round(time_to_first_token, 2),
                    "timestamp": datetime.now().isoformat()
                }
                session_manager.conversation_history.append(assistant_message_obj)
                
                if special_prompt_flag == 'save':
                    print("\n\n✅ Exporting content...")
                    filename = session_manager.export_filename
                    if not filename:
                        sanitized_topic = sanitize_filename(session_manager.topic_name)
                        extension = session_manager.export_type if session_manager.export_type else 'md'
                        filename = session_manager._get_unique_filename(session_manager.topic_name)
                        filename = os.path.join(session_manager.app_config.get("CHATS_DIRECTORY"), f"{os.path.basename(filename).split('.')[0]}.{extension}")
                    
                    content_to_save = extract_content_from_block(clean_assistant_response)
                    if content_to_save is None:
                        print(f"⚠️ Warning: Delimiters '{BLOCK_START}' and '{BLOCK_END}' not found. Saving the entire response as a fallback.")
                        content_to_save = clean_assistant_response
                    
                    save_content_to_file(content_to_save, filename, session_manager.export_type)

                session_manager.export_filename = None
                session_manager.export_type = None

                print(f"\n\nMetrics for this turn: User Tokens: {user_message_obj['token_count']} | Assistant Tokens: {assistant_message_obj['token_count']} | Response time: {total_time_for_turn:.2f}s | TTFT: {time_to_first_token:.2f}s")
            
            else:
                print("\n\n✅ Exporting entire conversation...")
                full_content_to_save = "\n\n".join([
                    f"[{msg['role'].upper()}] ({msg.get('timestamp', 'N/A')}) - Model: {msg.get('model_used', 'N/A')}:\n{msg['content']}"
                    for msg in session_manager.conversation_history if msg['role'] in ['user', 'assistant']
                ])
                
                filename = session_manager.export_filename
                if not filename:
                    sanitized_topic = sanitize_filename(session_manager.topic_name)
                    extension = session_manager.export_type if session_manager.export_type else 'md'
                    filename = session_manager._get_unique_filename(session_manager.topic_name)
                    filename = os.path.join(session_manager.app_config.get("CHATS_DIRECTORY"), f"{os.path.basename(filename).split('.')[0]}.{extension}")

                save_content_to_file(full_content_to_save, filename, session_manager.export_type)
                session_manager.export_filename = None
                session_manager.export_type = None
            
    except KeyboardInterrupt:
        print("\nEnding chat.")

    if session_manager.conversation_history:
        print("\n---------------------------\n")
        save_choice = input("Would you like to save this chat history? (y/n, default 'y'): ").strip().lower()
        if save_choice != 'n':
            session_manager.chat_data["total_response_time_seconds"] = round(total_response_time, 2)
            session_manager.chat_data["total_ttft_seconds"] = round(total_ttft, 2)
            session_manager.save_session()
        else:
            print("Chat history not saved.")
    else:
        print("No conversation to save.")

    print(f"\n--- Session Summary ({session_manager.model_name}) ---")
    print(f"Total model generation time: {total_response_time:.2f} seconds.")
    print(f"Total time to first token: {total_ttft:.2f} seconds.")
    print("Goodbye! 👋")

def clear_screen():
    """Clears the console screen."""
    os.system('cls' if os.name == 'nt' else 'clear')

# --- Functions for the Ollama Models Menu ---
def list_models(ollama_service: OllamaService) -> List[Tuple[str, int, str]]:
    """
    Lists available models, their size, and returns a sorted list of tuples for display.
    
    Args:
        ollama_service (OllamaService): The Ollama service instance.
        
    Returns:
        List[Tuple[str, int, str]]: A sorted list of model information tuples (name, size_bytes, size_str).
    """
    spinner = Halo(text='Fetching available models...', spinner='dots')
    spinner.start()
    try:
        response_obj = ollama_service.client.list()
        
        if not hasattr(response_obj, 'models'):
            spinner.fail("❌ Failed to fetch models: 'models' attribute not found in response.")
            print("Please ensure your Ollama server is running and accessible.")
            return []

        models = response_obj.models
        if not models:
            spinner.fail("❌ Ollama server returned an empty list of models.")
            return []

        model_list = []
        for model_info in models:
            name = model_info.get('model', 'Unknown')
            size_bytes = model_info.get('size', 0)
            size_mb = size_bytes / (1024 * 1024)
            size_str = f"{size_mb:.2f} MB" if size_bytes < 1024 * 1024 * 1024 else f"{size_bytes / (1024 * 1024 * 1024):.2f} GB"
            
            model_list.append((name, size_bytes, size_str))

        model_list.sort(key=lambda x: (x[1], x[0]))

        print('\n')
        for i, (name, _, size_str) in enumerate(model_list):
            print(f"[{i + 1}] {name} ({size_str})")
        spinner.succeed("✅ Successfully fetched model list.")
        return model_list

    except ollama.OllamaAPIError as e:
        spinner.fail(f"❌ Ollama API Error: {e}")
        print("Please check that the Ollama server is running and accessible at the configured host.")
        return []
    except Exception as e:
        spinner.fail(f"❌ An unexpected error occurred: {e}")
        return []

def delete_model(ollama_service: OllamaService):
    """
    Lists models and allows the user to delete one.
    
    Args:
        ollama_service (OllamaService): The Ollama service instance.
    """
    while True:
        clear_screen()
       
        available_models_info = list_models(ollama_service)
        clear_screen()
        
        print("--- Delete an Ollama Model ---")
        
        if not available_models_info:
            print("No models available to delete.")
            return

        print("\n--- Available Ollama Models ---")
        for i, (name, _, size_str) in enumerate(available_models_info):
            print(f"[{i + 1}] {name} ({size_str})")
        print("[0] Return to Main Menu")
        print("-------------------------------")

        try:
            choice = input("Enter the number of the model to delete: ").strip()
            if choice == '0':
                return

            model_index = int(choice) - 1
            if 0 <= model_index < len(available_models_info):
                model_to_delete = available_models_info[model_index][0]
                confirmation = input(f"Are you sure you want to delete '{model_to_delete}'? This cannot be undone. (y/n): ").strip().lower()
                
                if confirmation == 'y':
                    spinner = Halo(text=f'Deleting {model_to_delete}...', spinner='dots')
                    spinner.start()
                    try:
                        ollama_service.client.delete(model=model_to_delete)
                        spinner.succeed(f"✅ Model '{model_to_delete}' successfully deleted.")
                        
                        available_models_info.pop(model_index)
                        
                    except ollama.OllamaAPIError as e:
                        spinner.fail(f"❌ Failed to delete model: {e}")
                    except Exception as e:
                        spinner.fail(f"❌ An unexpected error occurred: {e}")
                else:
                    print("Deletion cancelled.")
                
                input("\nPress Enter to continue...")

            else:
                print("Invalid number. Please try again.")
                input("\nPress Enter to continue...")
        except ValueError:
            print("Invalid input. Please enter a number.")
            input("\nPress Enter to continue...")
            
def install_model(ollama_service: OllamaService):
    """
    Prompts the user to install a new model from the Ollama registry and shows progress.
    
    Args:
        ollama_service (OllamaService): The Ollama service instance.
    """
    clear_screen()
    print("--- Install an Ollama Model ---")
    print("If you're not sure which model to install, you can browse the official Ollama library here:")
    print("https://ollama.com/library")
    print("-" * 30)
    
    model_name = input("Enter the name of the model to install (e.g., 'llama3'): ").strip()

    if not model_name:
        print("Model name cannot be empty. Returning to menu.")
        input("\nPress Enter to continue...")
        return

    print(f"\nPulling model '{model_name}'. This may take some time depending on your connection.")

    try:
        stream = ollama_service.client.pull(model=model_name, stream=True)
        for progress in stream:
            status = progress.get('status', '')
            
            completed = progress.get('completed')
            total = progress.get('total')

            if completed is not None and total is not None and total > 0:
                percent = (completed / total) * 100
                bar_length = 40
                filled_length = int(bar_length * completed / total)
                bar = '█' * filled_length + '-' * (bar_length - filled_length)
                
                output_line = f"Status: {status: <20} [{bar}] {percent: >5.1f}%"
                print(output_line, end='\r', flush=True)
            else:
                output_line = f"Status: {status}"
                print(output_line, end='\r', flush=True)

        print("\n✅ Model downloaded successfully!")
    
    except Exception as e:
        print(f"\n❌ An error occurred: {e}")
        print("Please check that the Ollama server is running and the model name is correct.")

    input("\nPress Enter to continue...")

def models_menu(ollama_service: OllamaService):
    """The submenu for managing Ollama models."""
    while True:
        clear_screen()
        print("\n--- Ollama Models Menu ---")
        print("[1] List installed models")
        print("[2] Delete an installed model")
        print("[3] Install a model")
        print("[0] Return to Main Menu")
        choice = input("Enter your choice: ").strip().lower()

        if choice == '1':
            list_models(ollama_service)
            input("Press any key to continue ... ")
        elif choice == '2':
            delete_model(ollama_service)
        elif choice == '3':
            install_model(ollama_service)
        elif choice == '0':
            break
        else:
            print("Invalid choice. Please try again.")
        
def main_menu():
    """Main entry point for the application. Presents a menu for session management."""
    parser = argparse.ArgumentParser(description="Ollama Chat with Session Management.")
    parser.add_argument('--ingest', type=str, default=None, help='Path to a file to ingest as initial context.')
    parser.add_argument('--ollama-host', type=str, default=None, help='Override the Ollama host URL (e.g., http://192.168.1.100:11434).')
    args = parser.parse_args()
    
    app_config = AppConfig()
    
    if args.ollama_host:
        app_config.config['OLLAMA_HOST'] = args.ollama_host
    
    ollama_service = OllamaService(host=app_config.get("OLLAMA_HOST"))
    
    os.makedirs(app_config.get("CHATS_DIRECTORY"), exist_ok=True)
    if tokenizer:
        print("Token counting enabled for context management.")
    else:
        print("Token counting disabled. Context management will be based on message count.")
    
    while True:
        clear_screen()
        print("\n--- Main Menu ---")
        print("[1] Start a New Chat")
        print("[2] Load an Existing Chat")
        print("[3] Delete a Chat Session")
        print("[4] Search Chat Sessions")
        print("[5] Manage Ollama Models")
        print("[q] Quit")
        choice = input("Enter your choice: ").strip().lower()

        if choice == '1':
            available_models = ollama_service.get_available_models()
            if not available_models:
                print("No models found. Please pull a model using 'ollama pull <model_name>' and try again.")
                input("\nPress Enter to return to the main menu...")
                continue
            
            print("\n--- Available Models ---")
            for i, model in enumerate(available_models):
                print(f"[{i + 1}] {model}")
            
            model_choice_input = input(f"Enter the number of the model to use (or press Enter for '{app_config.get('DEFAULT_MODEL')}'): ")
            
            selected_model = app_config.get('DEFAULT_MODEL')
            try:
                model_index = int(model_choice_input) - 1
                if 0 <= model_index < len(available_models):
                    selected_model = available_models[model_index]
                else:
                    print(f"Invalid choice. Using default model: {app_config.get('DEFAULT_MODEL')}.")
            except (ValueError, IndexError):
                print(f"Invalid input. Using default model: {app_config.get('DEFAULT_MODEL')}.")
            
            new_topic = input("Enter a topic for the new chat (or press Enter for 'New Session'): ")
            sys_prompt_input = input(f"Enter a system prompt (or press Enter for '{app_config.get('DEFAULT_SYSTEM_PROMPT')}'): ")
            initial_prompt = sys_prompt_input if sys_prompt_input else app_config.get('DEFAULT_SYSTEM_PROMPT')
            
            session_manager = SessionManager(ollama_service, app_config, selected_model, system_prompt=initial_prompt, new_topic=new_topic if new_topic else None)
            run_chat_session(ollama_service, session_manager, app_config, args.ingest)
            
        elif choice == '2':
            clear_screen()
            files = list_sessions(app_config.get("CHATS_DIRECTORY"))
            if files:
                while True:
                    try:
                        session_input = input("Enter the number of the session to load (or '0' to go back'): ")
                        if session_input == '0':
                            break
                        
                        session_num = int(session_input)
                        if 1 <= session_num <= len(files):
                            file_path = os.path.join(app_config.get("CHATS_DIRECTORY"), files[session_num - 1])
                            with open(file_path, 'r', encoding='utf-8') as f:
                                loaded_data = json.load(f)
                                loaded_prompt = loaded_data.get('system_prompt', app_config.get('DEFAULT_SYSTEM_PROMPT'))
                            
                            clear_screen()
                            session_manager = SessionManager(
                                ollama_service=ollama_service, 
                                app_config=app_config, 
                                model_name=loaded_data.get('model_name', app_config.get('DEFAULT_MODEL')), 
                                system_prompt=loaded_prompt, 
                                session_file=files[session_num - 1]
                            )
                            run_chat_session(ollama_service, session_manager, app_config, args.ingest)
                            break
                        else:
                            print("Invalid session number.")
                    except ValueError:
                        print("Invalid input. Please enter a number.")
                        
        elif choice == '3':
            clear_screen()
            files = list_sessions(app_config.get("CHATS_DIRECTORY"))
            if files:
                while True:
                    try:
                        session_input = input("Enter the number of the session to delete (or '0' to go back'): ")
                        if session_input == '0':
                            break
                        session_num = int(session_input)
                        if 1 <= session_num <= len(files):
                            file_to_delete = os.path.join(app_config.get("CHATS_DIRECTORY"), files[session_num - 1])
                            confirm = input(f"Are you sure you want to delete '{files[session_num - 1]}'? (y/n): ").lower()
                            if confirm == 'y':
                                os.remove(file_to_delete)
                                print(f"Session '{files[session_num - 1]}' deleted.")
                                files.pop(session_num - 1)
                            else:
                                print("Deletion cancelled.")
                            continue
                        else:
                            print("Invalid session number.")
                    except (ValueError, FileNotFoundError, IndexError):
                        print("Invalid input or file not found.")
        elif choice == '4':
            while True:
                clear_screen()
                print("--- Search Chat Sessions ---")
                query_string = input("Enter your search query (with optional flags, e.g., 'hello --case-sensitive --role user'): ")
                if query_string.strip():
                    search_sessions(query_string, app_config.get("CHATS_DIRECTORY"))
                else:
                    print("Search query cannot be empty.")
                
                again = input("\nDo you want to perform another search? (y/n, default 'n'): ").strip().lower()
                if again != 'y':
                    break
        elif choice == '5':
            models_menu(ollama_service)
        elif choice == 'q':
            print("Exiting application. Goodbye! 👋")
            break
        else:
            print("Invalid choice. Please try again.")
            
        if choice in ['1', '2', '3', '4', '5']:
            input("\nPress Enter to return to the main menu...")

if __name__ == "__main__":
    main_menu()